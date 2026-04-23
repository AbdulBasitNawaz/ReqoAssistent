import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import google.generativeai as genai

# ---------------------- Streamlit Page Config ----------------------
st.set_page_config(page_title="Reqo Assistant", page_icon="Images/logo3.png", layout="wide")
st.title("📄 Generate Requirement Document")
st.write("You will generate a clean .docx requirements document and a concise abstract for further meatings, and will be forwarded to the team automatically when submitted.")

# ---------------------- Check Form Data ----------------------
if "form_data" not in st.session_state:
    st.warning("⚠ No form data found! Please fill the requirement form first.")
    st.stop()

data = st.session_state.form_data

# ---------------------- Configure Gemini ----------------------
genai.configure(api_key="AIzaSyDs_znLMGrtsksXdaHLI8owhdEhzIiifJo")
system_instruction = """
You are REQO — an AI Requirement Summarizer.
Generate a concise abstract suitable for display or download.
Keep it professional, 5-6 sentences maximum.
"""
model = genai.GenerativeModel(
    "gemini-2.5-flash",
    system_instruction=system_instruction
)

# ---------------------- Generate Abstract ----------------------
if "email_abstract" not in st.session_state:
    prompt = f"Create a short, professional abstract from the following requirement data:\n\n{data}"
    response = model.generate_content(prompt)
    st.session_state.email_abstract = response.text.strip()

# ---------------------- Show Abstract ----------------------
st.subheader("Project Abstract")
st.text_area("Abstract", st.session_state.email_abstract, height=200)

# ---------------------- Create DOCX ----------------------
doc = Document()
title = doc.add_heading("REQUIREMENT DOCUMENT", level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("Generated via REQO — Requirement Engineering Assistant")
doc.add_paragraph("")

def add_text(label, content):
    doc.add_heading(label, level=3)
    doc.add_paragraph(content if content else "N/A")
    doc.add_paragraph("")

def add_bullet_list(label, items):
    doc.add_heading(label, level=3)
    if not items:
        doc.add_paragraph("N/A")
    else:
        for item in items:
            doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("")

# Main Sections
add_text("Project Purpose", data.get("purpose"))
add_text("Problem Statement", data.get("problem"))
add_text("Project Vision", data.get("vision"))
add_bullet_list("Target Users", data.get("target_users", []))
add_text("User Goals", data.get("user_goals"))
add_bullet_list("Core Features", data.get("core_features", []))
add_text("Special Features", data.get("special_features"))
add_text("Platform", data.get("platform"))
add_text("Technology Preference", data.get("preference"))
add_text("Timeline", data.get("timeline"))
add_text("Budget", data.get("budget"))
add_text("Performance Requirements", data.get("performance"))
add_text("Security Requirements", data.get("security"))
add_text("Expected Users", data.get("expected_users"))
add_text("Client Name", data.get("name"))
add_text("Client Email", data.get("email"))
add_text("Client Phone", data.get("phone"))
add_bullet_list("Functional Requirements", data.get("functional_requirements", []))
add_bullet_list("Non-Functional Requirements", data.get("non_functional_requirements", []))
doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
doc.add_paragraph("Generated using REQO Assistant Prototype.")

# Save DOCX to buffer
buffer = BytesIO()
doc.save(buffer)
buffer.seek(0)

# ---------------------- Download Buttons ----------------------

col1, col2 = st.columns(2)

with col1:
    st.download_button(
        label="Download Abstract",
        data=st.session_state.email_abstract,
        file_name="project_abstract.txt",
        mime="text/plain"
    )

with col2:
    st.download_button(
        label="Download Requirement Document",
        data=buffer,
        file_name="requirement_document.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ---------------------- Navigation Buttons ----------------------

# Add custom CSS for green buttons
# ---------------------- Custom CSS for ALL buttons ----------------------

st.markdown(
    """
    <style>
    /* Style all Streamlit buttons including download buttons */
    div.stButton > button, div.stDownloadButton > button {
        background-color: #28a745;
        color: white;
        height: 3em;
        width: 100%;
        border-radius: 5px;
        border: none;
        font-size: 16px;
        font-weight: 600;
        transition: all 0.3s ease;
    }

    div.stButton > button:hover, div.stDownloadButton > button:hover {
        background-color: #218838;
        color: white;
        cursor: pointer;
    }
    </style>
    """,
    unsafe_allow_html=True
)

col1, col2 = st.columns(2)

with col1:
    if st.button("Submit Requirements"):
        st.success("Your requirements have been submitted! (Prototype only)")

with col2:
    if st.button("Go to Chatbot"):
        st.switch_page("pages/chat.py")

